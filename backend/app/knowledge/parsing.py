"""Document parsing (Phase 2.1): pdf/docx/md/txt → normalized ParsedDocument.

Deterministic, framework-light, untrusted-input-safe. Runs synchronously;
callers execute it via ``asyncio.to_thread`` and bound it with a timeout at
the service boundary (no process/container isolation — async-boundary only).

Guards: size/type are enforced upstream (DocumentService); page count,
DOCX archive/uncompressed-size limits, empty-extraction detection, and
timeout are enforced here or at the service boundary.

Parsed text is an in-memory handoff to Phase 2.2 ingestion — never persisted
(Phase 1 schema deliberately has no parsed_text column).
"""

from __future__ import annotations

import asyncio
import zipfile
from dataclasses import dataclass
from io import BytesIO

from docx import Document as DocxDocument
from markdown_it import MarkdownIt
from pypdf import PdfReader

from app.domain.enums import DocumentKind
from app.domain.errors import ValidationFailedError

# DOCX is a zip archive: bound the total uncompressed size to blunt
# decompression-bomb resource exhaustion (upload cap is 5 MB; conservative
# default is 50x that for the extracted archive contents).
MAX_DOCX_UNCOMPRESSED_BYTES = 5 * 50 * 1024 * 1024

# MIME constants mirroring DocumentService._ALLOWED_MIME (keep in sync).
MIME_PDF = "application/pdf"
MIME_DOCX = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
MIME_MARKDOWN = "text/markdown"
MIME_PLAIN = "text/plain"


@dataclass(frozen=True)
class ParsedDocument:
    """Normalized output of document parsing (in-memory handoff to 2.2)."""

    content: str
    # Canonical format label: "pdf" | "docx" | "markdown" | "text"
    format: str
    page_count: int
    # Source metadata (mirrors the persisted Document row where applicable).
    filename: str
    mime: str
    size: int
    content_hash: str


def parse_document(
    *,
    data: bytes,
    kind: DocumentKind,
    mime: str,
    filename: str,
    content_hash: str,
    max_pages: int,
) -> ParsedDocument:
    """Parse one document into a normalized representation.

    Raises ValidationFailedError with actionable details on malformed files,
    empty extraction, page-limit breach, or unsupported type.
    """
    if not data:
        raise ValidationFailedError("empty document")
    if mime == MIME_PDF:
        return _parse_pdf(data, kind, filename, content_hash, max_pages)
    if mime == MIME_DOCX:
        return _parse_docx(data, kind, filename, content_hash)
    if mime == MIME_MARKDOWN:
        return _parse_markdown(data, kind, filename, content_hash)
    if mime == MIME_PLAIN:
        return _parse_text(data, kind, filename, content_hash)
    raise ValidationFailedError(
        f"unsupported document type: {mime}", details={"mime": mime, "filename": filename}
    )


async def parse_document_with_timeout(
    *,
    data: bytes,
    kind: DocumentKind,
    mime: str,
    filename: str,
    content_hash: str,
    max_pages: int,
    timeout_seconds: float,
) -> ParsedDocument:
    """Run the synchronous parser off the event loop, bounded by a timeout.

    NOTE: ``asyncio.to_thread`` provides async-boundary isolation only — it
    keeps blocking parser work off the event loop. It is not process-level
    isolation; true isolation remains a later hardening concern.
    """
    return await asyncio.wait_for(
        asyncio.to_thread(
            parse_document,
            data=data,
            kind=kind,
            mime=mime,
            filename=filename,
            content_hash=content_hash,
            max_pages=max_pages,
        ),
        timeout=timeout_seconds,
    )


def _parse_pdf(
    data: bytes, kind: DocumentKind, filename: str, content_hash: str, max_pages: int
) -> ParsedDocument:
    try:
        reader = PdfReader(BytesIO(data))
        page_count = len(reader.pages)
        if page_count > max_pages:
            raise ValidationFailedError(
                f"document exceeds page limit ({page_count} > {max_pages})",
                details={"pages": page_count, "max_pages": max_pages, "filename": filename},
            )
        parts: list[str] = []
        for page in reader.pages:
            text = page.extract_text() or ""
            if text.strip():
                parts.append(text.strip())
        content = "\n\n".join(parts)
    except ValidationFailedError:
        raise
    except Exception as exc:  # pypdf raises assorted exceptions on malformed input
        raise ValidationFailedError(
            f"failed to parse PDF: {exc}", details={"filename": filename, "kind": kind.value}
        ) from exc
    _require_content(content, filename, kind)
    return ParsedDocument(
        content=content,
        format="pdf",
        page_count=page_count,
        filename=filename,
        mime=MIME_PDF,
        size=len(data),
        content_hash=content_hash,
    )


def _parse_docx(
    data: bytes, kind: DocumentKind, filename: str, content_hash: str
) -> ParsedDocument:
    try:
        with zipfile.ZipFile(BytesIO(data)) as zf:
            total = sum(info.file_size for info in zf.infolist())
            if total > MAX_DOCX_UNCOMPRESSED_BYTES:
                raise ValidationFailedError(
                    "DOCX archive exceeds uncompressed size limit",
                    details={"uncompressed_bytes": total, "max_bytes": MAX_DOCX_UNCOMPRESSED_BYTES},
                )
            doc = DocxDocument(BytesIO(data))
    except ValidationFailedError:
        raise
    except Exception as exc:  # zipfile/docx raise assorted exceptions on malformed input
        raise ValidationFailedError(
            f"failed to parse DOCX: {exc}", details={"filename": filename, "kind": kind.value}
        ) from exc
    parts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    content = "\n\n".join(parts)
    _require_content(content, filename, kind)
    return ParsedDocument(
        content=content,
        format="docx",
        page_count=1,
        filename=filename,
        mime=MIME_DOCX,
        size=len(data),
        content_hash=content_hash,
    )


def _parse_markdown(
    data: bytes, kind: DocumentKind, filename: str, content_hash: str
) -> ParsedDocument:
    try:
        text = data.decode("utf-8")
    except UnicodeDecodeError:
        text = data.decode("latin-1")
    # markdown-it gives a deterministic token stream; plain text extraction
    # keeps inline text (incl. emphasis/links) and code blocks.
    md = MarkdownIt("commonmark", {"html": False})
    parts: list[str] = []
    for token in md.parse(text):
        if token.type == "inline" and token.children:
            inline = "".join(
                child.content for child in token.children if child.type in ("text", "code_inline")
            )
            if inline.strip():
                parts.append(inline.strip())
        elif token.type in ("code_block", "fence") and token.content.strip():
            parts.append(token.content.strip())
    content = "\n\n".join(parts)
    _require_content(content, filename, kind)
    return ParsedDocument(
        content=content,
        format="markdown",
        page_count=1,
        filename=filename,
        mime=MIME_MARKDOWN,
        size=len(data),
        content_hash=content_hash,
    )


def _parse_text(
    data: bytes, kind: DocumentKind, filename: str, content_hash: str
) -> ParsedDocument:
    try:
        content = data.decode("utf-8").strip()
    except UnicodeDecodeError:
        content = data.decode("latin-1").strip()
    _require_content(content, filename, kind)
    return ParsedDocument(
        content=content,
        format="text",
        page_count=1,
        filename=filename,
        mime=MIME_PLAIN,
        size=len(data),
        content_hash=content_hash,
    )


def _require_content(content: str, filename: str, kind: DocumentKind) -> None:
    if not content.strip():
        raise ValidationFailedError(
            "no extractable text in document",
            details={"filename": filename, "kind": kind.value},
        )
